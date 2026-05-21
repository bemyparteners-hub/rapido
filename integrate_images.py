#!/usr/bin/env python3
"""
integrate_images.py
Insère automatiquement les images PNG dans les cellules-placeholders
du document Word Note_3_Modes_Constructifs_PAPILLON.
"""

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

INPUT_DOC = "Note_3_Modes_Constructifs_PAPILLON (1).docx"
OUTPUT_DOC = "Note_3_Modes_Constructifs_PAPILLON_avec_images.docx"
IMAGES_DIR = Path("images")

# Préfixe fichier → (libellé dans placeholder, nb étapes)
MODE_MAPPING = {
    "Voile banché":               ("voile_", 10),
    "Dalle pleine":               ("dalle_",  9),
    "Plancher Cofraplus 60":      ("cofra_",  9),
    "Fondations superficielles":  ("fond_",   9),
    "Micropieux IGU":             ("micro_",  9),
    "RSO par plots":              ("rso_",   10),
}

# Lookup inverse : préfixe → libellé (pour rapport des fichiers non utilisés)
PREFIX_TO_MODE = {v[0]: k for k, v in MODE_MAPPING.items()}

# Regex de détection d'un placeholder
# Gère apostrophe droite ' (U+0027) et courbe ' (U+2019), tiret cadratin — (U+2014) et -
PATTERN = re.compile(
    r"\[\s*INS[EÉ]RER ICI L['’]IMAGE\s*[—\-]\s*(.+?)\s*[—\-]\s*[EÉ]tape\s*(\d+)\s*:\s*(.+?)\s*\]",
    re.IGNORECASE,
)

IMAGE_WIDTH = Cm(11.9)

# ---------------------------------------------------------------------------
# Helpers XML
# ---------------------------------------------------------------------------

def cell_has_image(cell) -> bool:
    """Renvoie True si la cellule contient déjà un élément image (Drawing/Pict)."""
    for el in cell._tc.iter():
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag in ("drawing", "pict"):
            return True
    return False


def clear_cell(cell) -> None:
    """Vide le contenu d'une cellule en conservant ses propriétés (shading, bordures)."""
    tc = cell._tc
    tcPr_tag = qn("w:tcPr")
    for child in list(tc):
        if child.tag != tcPr_tag:
            tc.remove(child)
    # Les cellules Word doivent toujours contenir au moins un paragraphe
    tc.append(OxmlElement("w:p"))


def set_paragraph_spacing(para, before_pt: float, after_pt: float) -> None:
    pPr = para._p.get_or_add_pPr()
    # Chercher ou créer <w:spacing>
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    # Valeurs en twips (1 pt = 20 twips)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"),  str(int(after_pt  * 20)))


# ---------------------------------------------------------------------------
# Insertion d'une image dans une cellule
# ---------------------------------------------------------------------------

def insert_image_in_cell(cell, img_path: Path, mode_short: str,
                          step_num: str, step_title: str, errors: list) -> bool:
    """
    Remplace le contenu de la cellule par l'image + légende.
    Renvoie True si réussi, False en cas d'erreur.
    """
    try:
        from PIL import Image as PILImage
        PILImage.open(img_path).verify()          # valide que le fichier n'est pas corrompu
    except Exception as exc:
        errors.append(f"Image corrompue ({img_path.name}): {exc}")
        return False

    clear_cell(cell)

    # --- Paragraphe image centré ---
    para_img = cell.paragraphs[0]
    para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para_img, before_pt=10, after_pt=3)
    run_img = para_img.add_run()
    run_img.add_picture(str(img_path), width=IMAGE_WIDTH)

    # --- Paragraphe légende centré ---
    p_elem = OxmlElement("w:p")
    cell._tc.append(p_elem)
    para_cap = cell.paragraphs[1]
    para_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para_cap, before_pt=0, after_pt=12)

    caption_text = f"Figure — {mode_short} — Étape {step_num} : {step_title}"
    run_cap = para_cap.add_run(caption_text)
    run_cap.font.name = "Calibri"
    run_cap.font.size = Pt(9)
    run_cap.font.italic = True
    run_cap.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    return True


# ---------------------------------------------------------------------------
# Scan récursif des tableaux
# ---------------------------------------------------------------------------

def scan_tables(tables, doc, dry_run: bool, force: bool,
                stats: dict, errors: list, used_images: set) -> None:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                # Scan des tableaux imbriqués d'abord
                if cell.tables:
                    scan_tables(cell.tables, doc, dry_run, force,
                                stats, errors, used_images)

                # Cellule déjà illustrée ?
                if cell_has_image(cell):
                    if not force:
                        stats["already_illustrated"] += 1
                        continue
                    # --force : on re-traite quand même, mais seulement si c'est
                    # une cellule dont le texte *après* vide ressemble à un placeholder
                    # (cas impossible si c'est déjà une image sans texte) → skip
                    stats["already_illustrated"] += 1
                    continue

                text = cell.text.strip()
                if not text:
                    continue

                m = PATTERN.search(text)
                if not m:
                    continue

                mode_short = m.group(1).strip()
                step_num   = m.group(2).strip()
                step_title = m.group(3).strip()

                if mode_short not in MODE_MAPPING:
                    errors.append(
                        f"Mode inconnu dans placeholder : {mode_short!r} "
                        f"(Étape {step_num})"
                    )
                    stats["placeholders_skipped"] += 1
                    continue

                prefix, _ = MODE_MAPPING[mode_short]
                img_path = IMAGES_DIR / f"{prefix}{step_num}.png"

                stats["modes"].setdefault(mode_short, {"integrated": 0, "missing": []})

                if not img_path.exists():
                    stats["modes"][mode_short]["missing"].append(img_path.name)
                    stats["placeholders_skipped"] += 1
                    continue

                used_images.add(img_path.name)

                if dry_run:
                    print(f"  [DRY-RUN] Insérerait {img_path.name} "
                          f"→ {mode_short} Étape {step_num}: {step_title}")
                    stats["modes"][mode_short]["integrated"] += 1
                    stats["integrated_total"] += 1
                else:
                    ok = insert_image_in_cell(
                        cell, img_path, mode_short, step_num, step_title, errors
                    )
                    if ok:
                        stats["modes"][mode_short]["integrated"] += 1
                        stats["integrated_total"] += 1
                    else:
                        stats["placeholders_skipped"] += 1


# ---------------------------------------------------------------------------
# Rapport
# ---------------------------------------------------------------------------

def print_report(stats: dict, used_images: set, errors: list,
                 dry_run: bool, output_path: str) -> None:
    label = " (DRY-RUN)" if dry_run else ""
    print()
    print("=" * 44)
    print(f"INTÉGRATION DES IMAGES — RAPPORT{label}")
    print("=" * 44)
    print()

    max_expected = {
        "Voile banché":              10,
        "Dalle pleine":               9,
        "Plancher Cofraplus 60":      9,
        "Fondations superficielles":  9,
        "Micropieux IGU":             9,
        "RSO par plots":             10,
    }

    for mode, (prefix, total) in MODE_MAPPING.items():
        m_stats = stats["modes"].get(mode, {"integrated": 0, "missing": []})
        integrated = m_stats["integrated"]
        missing    = m_stats["missing"]
        missing_str = (f"  (manquantes : {', '.join(missing)})" if missing else "")
        print(f"  {mode:<30}:  {integrated}/{total}  intégrées{missing_str}")

    total_expected = sum(v for _, (_, v) in MODE_MAPPING.items())
    print()
    print("-" * 44)
    print(f"  {'Total':<30}:  {stats['integrated_total']}/{total_expected}  intégrées")
    print(f"  Placeholders restants        :  {stats['placeholders_skipped']}")
    print(f"  Cellules déjà illustrées     :  {stats['already_illustrated']}")

    # Fichiers présents mais non utilisés
    all_present = {p.name for p in IMAGES_DIR.glob("*.png")}
    unused = sorted(all_present - used_images)
    print()
    print("  Fichiers image présents non utilisés (aucun placeholder correspondant) :")
    if unused:
        for f in unused:
            print(f"    - {f}")
    else:
        print("    aucun")

    print()
    print("  Erreurs rencontrées :")
    if errors:
        for e in errors:
            print(f"    - {e}")
    else:
        print("    aucune")

    print()
    if not dry_run:
        print(f"  Document sauvegardé : {output_path}")
    print("=" * 44)
    print()


# ---------------------------------------------------------------------------
# Point d'entrée
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Intègre les images PNG dans les placeholders du document Word."
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Liste les actions sans modifier le document."
    )
    parser.add_argument(
        "--force", action="store_true",
        help="Re-insère même si une image est déjà présente dans la cellule."
    )
    args = parser.parse_args()

    if not Path(INPUT_DOC).exists():
        sys.exit(f"Erreur : document introuvable → {INPUT_DOC}")
    if not IMAGES_DIR.is_dir():
        sys.exit(f"Erreur : dossier images introuvable → {IMAGES_DIR}")

    print(f"Ouverture de : {INPUT_DOC}")
    doc = Document(INPUT_DOC)

    stats: dict = {
        "integrated_total": 0,
        "placeholders_skipped": 0,
        "already_illustrated": 0,
        "modes": {},
    }
    errors: list = []
    used_images: set = set()

    scan_tables(doc.tables, doc, args.dry_run, args.force,
                stats, errors, used_images)

    if not args.dry_run:
        print(f"Sauvegarde sous : {OUTPUT_DOC}")
        doc.save(OUTPUT_DOC)

    print_report(stats, used_images, errors, args.dry_run, OUTPUT_DOC)


if __name__ == "__main__":
    main()
